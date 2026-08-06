import docx
from docx.shared import Inches, Pt, Cm, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT, WD_ALIGN_VERTICAL
from docx.oxml import OxmlElement, parse_xml
from docx.oxml.ns import nsdecls, qn
import datetime
import re
import io
import zipfile

# --- HELPER FUNCTIONS FOR STANDARD FORMATTING ---

def set_a4_margins(section):
    """Sets A4 margins according to standard administrative formatting."""
    section.page_width = Cm(21.0)
    section.page_height = Cm(29.7)
    section.top_margin = Cm(2.0)
    section.bottom_margin = Cm(2.0)
    section.left_margin = Cm(3.0)
    section.right_margin = Cm(1.5)

def set_cell_border(cell, **kwargs):
    """
    Set cell's border
    kwargs: top, bottom, left, right
    values: dict(sz=12, val='single', color='000000', space='0')
    """
    tcPr = cell._tc.get_or_add_tcPr()
    tcBorders = parse_xml(r'<w:tcBorders %s/>' % nsdecls('w'))
    for edge in ('top', 'left', 'bottom', 'right', 'insideH', 'insideV'):
        edge_data = kwargs.get(edge)
        if edge_data:
            tag = 'w:{}'.format(edge)
            element = parse_xml(r'<%s %s w:val="%s" w:sz="%s" w:space="%s" w:color="%s"/>' % (
                tag, nsdecls('w'),
                edge_data.get('val', 'single'),
                edge_data.get('sz', '4'),
                edge_data.get('space', '0'),
                edge_data.get('color', '000000')
            ))
            tcBorders.append(element)
    tcPr.append(tcBorders)

def set_table_borders(table):
    """Applies clean single-line borders to all cells in table."""
    border_kwargs = {
        'top': {'sz': '4', 'val': 'single', 'color': '000000', 'space': '0'},
        'bottom': {'sz': '4', 'val': 'single', 'color': '000000', 'space': '0'},
        'left': {'sz': '4', 'val': 'single', 'color': '000000', 'space': '0'},
        'right': {'sz': '4', 'val': 'single', 'color': '000000', 'space': '0'},
    }
    for row in table.rows:
        for cell in row.cells:
            set_cell_border(cell, **border_kwargs)

def format_run(run, font_name="Times New Roman", font_size_pt=13, bold=False, italic=False, color=None):
    """Formats a run with font name, size, bold, italic, color."""
    run.font.name = font_name
    run.font.size = Pt(font_size_pt)
    run.bold = bold
    run.italic = italic
    if color:
        run.font.color.rgb = color

def set_paragraph_format(p, alignment=WD_ALIGN_PARAGRAPH.LEFT, space_before=0, space_after=3, line_spacing=1.15, first_line_indent=0):
    """Sets paragraph alignment, spacing and first line indent."""
    p.alignment = alignment
    p.paragraph_format.space_before = Pt(space_before)
    p.paragraph_format.space_after = Pt(space_after)
    p.paragraph_format.line_spacing = line_spacing
    if first_line_indent > 0:
        p.paragraph_format.first_line_indent = Pt(first_line_indent)

# --- PLACEHOLDER REPLACEMENT ENGINE ---

def replace_text_in_paragraph(paragraph, data_dict):
    """
    Replaces placeholders like {{KEY}} in paragraph text accurately.
    Handles split runs by evaluating full paragraph text when placeholders exist.
    """
    full_text = paragraph.text
    if "{{" not in full_text:
        return

    # Find all placeholder tags in text
    placeholders = re.findall(r"\{\{([^{}]+)\}\}", full_text)
    if not placeholders:
        return

    updated_text = full_text
    for tag in placeholders:
        key = tag.strip()
        val = str(data_dict.get(key, data_dict.get(tag, f"{{{{{tag}}}}}")))
        updated_text = updated_text.replace(f"{{{{{tag}}}}}", val)

    # Re-apply text preserving font if paragraph has runs
    if paragraph.runs:
        font_name = paragraph.runs[0].font.name or "Times New Roman"
        font_size = paragraph.runs[0].font.size or Pt(13)
        bold = paragraph.runs[0].bold
        italic = paragraph.runs[0].italic

        # Clear text in all runs except first
        paragraph.runs[0].text = updated_text
        for run in paragraph.runs[1:]:
            run.text = ""

        format_run(paragraph.runs[0], font_name=font_name, font_size_pt=font_size.pt if hasattr(font_size, 'pt') else 13, bold=bold, italic=italic)
    else:
        p_run = paragraph.add_run(updated_text)
        format_run(p_run, font_name="Times New Roman", font_size_pt=13)

def fill_docx_placeholders(doc, data_dict):
    """
    Fills placeholders {{KEY}} in doc paragraphs and tables.
    """
    # 1. Paragraphs in document body
    for paragraph in doc.paragraphs:
        replace_text_in_paragraph(paragraph, data_dict)

    # 2. Tables in document body
    for table in doc.tables:
        for row in table.rows:
            for cell in row.cells:
                for paragraph in cell.paragraphs:
                    replace_text_in_paragraph(paragraph, data_dict)

    # 3. Headers and footers
    for section in doc.sections:
        for paragraph in section.header.paragraphs:
            replace_text_in_paragraph(paragraph, data_dict)
        for paragraph in section.footer.paragraphs:
            replace_text_in_paragraph(paragraph, data_dict)

# --- TABLE APPEND ENGINE FOR NOTICE ---

def append_table_rows(doc, df_list):
    """
    Appends rows from df_list to the notice summary table.
    Expects table with header columns: STT, Họ và tên, Ngày sinh, Thôn, Dạng tật, Mức độ khuyết tật, Ghi chú.
    """
    if not doc.tables:
        return

    # Find the summary table (usually the main/last table or table with 'Họ và tên')
    target_table = None
    for tbl in doc.tables:
        header_text = "".join([c.text for c in tbl.rows[0].cells])
        if "Họ và tên" in header_text or "Họ và Tên" in header_text:
            target_table = tbl
            break

    if not target_table:
        target_table = doc.tables[0]

    # Remove template sample rows if present (rows after header)
    while len(target_table.rows) > 1:
        tr = target_table.rows[-1]._tr
        target_table._tbl.remove(tr)

    # Add rows from df_list
    for idx, row in enumerate(df_list, start=1):
        tr = target_table.add_row()
        tr.height = Pt(22)
        cells = tr.cells

        ho_ten = str(row.get('HO_TEN', ''))
        ngay_sinh = str(row.get('NGAY_SINH', ''))
        thon = str(row.get('THON', ''))
        dang_tat = str(row.get('DANG_TAT', ''))
        muc_do = str(row.get('MUC_DO', ''))
        ghi_chu = str(row.get('GHI_CHU', ''))

        values = [
            str(idx),
            ho_ten,
            ngay_sinh,
            thon,
            dang_tat,
            muc_do,
            ghi_chu
        ]

        alignments = [
            WD_ALIGN_PARAGRAPH.CENTER,  # STT
            WD_ALIGN_PARAGRAPH.LEFT,    # Họ tên
            WD_ALIGN_PARAGRAPH.CENTER,  # Ngày sinh
            WD_ALIGN_PARAGRAPH.LEFT,    # Thôn
            WD_ALIGN_PARAGRAPH.LEFT,    # Dạng tật
            WD_ALIGN_PARAGRAPH.CENTER,  # Mức độ
            WD_ALIGN_PARAGRAPH.LEFT     # Ghi chú
        ]

        for i, val in enumerate(values):
            if i < len(cells):
                cell = cells[i]
                cell.vertical_alignment = WD_ALIGN_VERTICAL.CENTER
                p = cell.paragraphs[0]
                p.text = val
                set_paragraph_format(p, alignment=alignments[i], space_before=2, space_after=2, line_spacing=1.0)
                if p.runs:
                    format_run(p.runs[0], font_name="Times New Roman", font_size_pt=12)

    set_table_borders(target_table)

# --- DEFAULT DOCUMENT CREATORS (Strict Standard Layouts) ---

def create_phieu_duoi_6_doc(data_dict=None):
    """Creates the 'Phiếu Xác Định Mức Độ Khuyết Tật Dưới 6 Tuổi' document according to Mau 02."""
    doc = docx.Document()
    section = doc.sections[0]
    set_a4_margins(section)

    # Header Table (Left: Agency, Right: Motto)
    table = doc.add_table(rows=1, cols=2)
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    table.autofit = False

    # Widths: Left 9cm, Right 10cm
    cell_l, cell_r = table.rows[0].cells
    cell_l.width = Cm(8.5)
    cell_r.width = Cm(9.5)

    p_l1 = cell_l.paragraphs[0]
    p_l1.text = "HỘI ĐỒNG XÁC ĐỊNH MĐKT"
    set_paragraph_format(p_l1, WD_ALIGN_PARAGRAPH.CENTER, space_after=1)
    format_run(p_l1.runs[0], font_name="Times New Roman", font_size_pt=12, bold=True)

    p_l2 = cell_l.add_paragraph("XÃ {{XA_UPPER}}")
    set_paragraph_format(p_l2, WD_ALIGN_PARAGRAPH.CENTER, space_after=1)
    format_run(p_l2.runs[0], font_name="Times New Roman", font_size_pt=12, bold=True)

    p_r1 = cell_r.paragraphs[0]
    p_r1.text = "CỘNG HÒA XÃ HỘI CHỦ NGHĨA VIỆT NAM"
    set_paragraph_format(p_r1, WD_ALIGN_PARAGRAPH.CENTER, space_after=1)
    format_run(p_r1.runs[0], font_name="Times New Roman", font_size_pt=12, bold=True)

    p_r2 = cell_r.add_paragraph("Độc lập - Tự do - Hạnh phúc")
    set_paragraph_format(p_r2, WD_ALIGN_PARAGRAPH.CENTER, space_after=6)
    format_run(p_r2.runs[0], font_name="Times New Roman", font_size_pt=13, bold=True)

    # Title
    p_title = doc.add_paragraph()
    p_title.text = "PHIẾU XÁC ĐỊNH MỨC ĐỘ KHUYẾT TẬT ĐỐI VỚI TRẺ EM DƯỚI 6 TUỔI"
    set_paragraph_format(p_title, WD_ALIGN_PARAGRAPH.CENTER, space_before=12, space_after=12)
    format_run(p_title.runs[0], font_name="Times New Roman", font_size_pt=14, bold=True)

    # Section I
    p_s1 = doc.add_paragraph("I. Thông tin người được xác định mức độ khuyết tật")
    set_paragraph_format(p_s1, WD_ALIGN_PARAGRAPH.LEFT, space_before=6, space_after=3)
    format_run(p_s1.runs[0], font_name="Times New Roman", font_size_pt=13, bold=True)

    lines_s1 = [
        "- Họ và tên: {{HO_TEN}}",
        "- Sinh ngày {{NGAY_SINH_NGAY}} tháng {{NGAY_SINH_THANG}} năm {{NGAY_SINH_NAM}}    Giới tính: {{GIOI_TINH}}",
        "- Hộ khẩu thường trú: Thôn {{THON}}, xã {{XA}}, huyện {{HUYEN}}, tỉnh {{TINH}}",
        "- Nơi ở hiện nay: Thôn {{THON}}, xã {{XA}}, huyện {{HUYEN}}, tỉnh {{TINH}}"
    ]
    for line in lines_s1:
        p = doc.add_paragraph(line)
        set_paragraph_format(p, space_after=2)
        format_run(p.runs[0], font_name="Times New Roman", font_size_pt=13)

    # Section II
    p_s2 = doc.add_paragraph("II. Thông tin người đại diện hợp pháp (nếu có)")
    set_paragraph_format(p_s2, WD_ALIGN_PARAGRAPH.LEFT, space_before=6, space_after=3)
    format_run(p_s2.runs[0], font_name="Times New Roman", font_size_pt=13, bold=True)

    lines_s2 = [
        "- Họ và tên người đại diện: {{HO_TEN_NDH}}",
        "- Mối quan hệ với người được xác định khuyết tật: {{MOI_QUAN_HE}}",
        "- Số CMND/CCCD: {{CMND_NDH}}",
        "- Hộ khẩu thường trú: Thôn {{THON}}, xã {{XA}}, huyện {{HUYEN}}, tỉnh {{TINH}}",
        "- Số điện thoại liên hệ: {{SDT_NDH}}"
    ]
    for line in lines_s2:
        p = doc.add_paragraph(line)
        set_paragraph_format(p, space_after=2)
        format_run(p.runs[0], font_name="Times New Roman", font_size_pt=13)

    # Section III
    p_s3 = doc.add_paragraph("III. Xác định dạng khuyết tật")
    set_paragraph_format(p_s3, WD_ALIGN_PARAGRAPH.LEFT, space_before=6, space_after=3)
    format_run(p_s3.runs[0], font_name="Times New Roman", font_size_pt=13, bold=True)

    p_s3_sub = doc.add_paragraph("Đã kiểm tra và đánh giá các dạng khuyết tật theo quy định ban hành kèm theo Thông tư số 01/2019/TT-BLĐTBXH đối với trẻ em dưới 6 tuổi.")
    set_paragraph_format(p_s3_sub, space_after=4)
    format_run(p_s3_sub.runs[0], font_name="Times New Roman", font_size_pt=13, italic=True)

    # Section IV
    p_s4 = doc.add_paragraph("IV. Xác định mức độ khuyết tật")
    set_paragraph_format(p_s4, WD_ALIGN_PARAGRAPH.LEFT, space_before=6, space_after=3)
    format_run(p_s4.runs[0], font_name="Times New Roman", font_size_pt=13, bold=True)

    # Section V
    p_s5 = doc.add_paragraph("V. Đề xuất kết luận dạng khuyết tật và mức độ khuyết tật")
    set_paragraph_format(p_s5, WD_ALIGN_PARAGRAPH.LEFT, space_before=6, space_after=3)
    format_run(p_s5.runs[0], font_name="Times New Roman", font_size_pt=13, bold=True)

    lines_s5 = [
        "1. Dạng khuyết tật: {{DANG_TAT}}",
        "2. Mức độ khuyết tật: {{MUC_DO}}",
        "3. Ghi chú bổ sung: {{GHI_CHU}}"
    ]
    for line in lines_s5:
        p = doc.add_paragraph(line)
        set_paragraph_format(p, space_after=3)
        format_run(p.runs[0], font_name="Times New Roman", font_size_pt=13)

    # Signatures
    p_date = doc.add_paragraph("{{XA}}, ngày {{NGAY_HOP_NGAY}} tháng {{NGAY_HOP_THANG}} năm {{NGAY_HOP_NAM}}")
    set_paragraph_format(p_date, WD_ALIGN_PARAGRAPH.RIGHT, space_before=12, space_after=4)
    format_run(p_date.runs[0], font_name="Times New Roman", font_size_pt=13, italic=True)

    sig_table = doc.add_table(rows=2, cols=2)
    sig_table.alignment = WD_TABLE_ALIGNMENT.CENTER
    c1, c2 = sig_table.rows[0].cells

    p_c1 = c1.paragraphs[0]
    p_c1.text = "NGƯỜI GHI PHIẾU\n(Ký, ghi rõ họ tên)"
    set_paragraph_format(p_c1, WD_ALIGN_PARAGRAPH.CENTER)
    format_run(p_c1.runs[0], font_name="Times New Roman", font_size_pt=12, bold=True)

    p_c2 = c2.paragraphs[0]
    p_c2.text = "CHỦ TỊCH HỘI ĐỒNG\n(Ký tên, đóng dấu)"
    set_paragraph_format(p_c2, WD_ALIGN_PARAGRAPH.CENTER)
    format_run(p_c2.runs[0], font_name="Times New Roman", font_size_pt=12, bold=True)

    c1_b, c2_b = sig_table.rows[1].cells
    p_c1_b = c1_b.paragraphs[0]
    p_c1_b.text = "\n\n\n{{NGUOI_LAP}}"
    set_paragraph_format(p_c1_b, WD_ALIGN_PARAGRAPH.CENTER)
    format_run(p_c1_b.runs[0], font_name="Times New Roman", font_size_pt=12, bold=True)

    p_c2_b = c2_b.paragraphs[0]
    p_c2_b.text = "\n\n\n{{CHU_TICH}}"
    set_paragraph_format(p_c2_b, WD_ALIGN_PARAGRAPH.CENTER)
    format_run(p_c2_b.runs[0], font_name="Times New Roman", font_size_pt=12, bold=True)

    if data_dict:
        fill_docx_placeholders(doc, data_dict)

    return doc


def create_phieu_tren_6_doc(data_dict=None):
    """Creates the 'Phiếu Xác Định Mức Độ Khuyết Tật Từ 6 Tuổi Trở Lên' document."""
    doc = docx.Document()
    section = doc.sections[0]
    set_a4_margins(section)

    # Header Table
    table = doc.add_table(rows=1, cols=2)
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    table.autofit = False

    cell_l, cell_r = table.rows[0].cells
    cell_l.width = Cm(8.5)
    cell_r.width = Cm(9.5)

    p_l1 = cell_l.paragraphs[0]
    p_l1.text = "HỘI ĐỒNG XÁC ĐỊNH MĐKT"
    set_paragraph_format(p_l1, WD_ALIGN_PARAGRAPH.CENTER, space_after=1)
    format_run(p_l1.runs[0], font_name="Times New Roman", font_size_pt=12, bold=True)

    p_l2 = cell_l.add_paragraph("XÃ {{XA_UPPER}}")
    set_paragraph_format(p_l2, WD_ALIGN_PARAGRAPH.CENTER, space_after=1)
    format_run(p_l2.runs[0], font_name="Times New Roman", font_size_pt=12, bold=True)

    p_r1 = cell_r.paragraphs[0]
    p_r1.text = "CỘNG HÒA XÃ HỘI CHỦ NGHĨA VIỆT NAM"
    set_paragraph_format(p_r1, WD_ALIGN_PARAGRAPH.CENTER, space_after=1)
    format_run(p_r1.runs[0], font_name="Times New Roman", font_size_pt=12, bold=True)

    p_r2 = cell_r.add_paragraph("Độc lập - Tự do - Hạnh phúc")
    set_paragraph_format(p_r2, WD_ALIGN_PARAGRAPH.CENTER, space_after=6)
    format_run(p_r2.runs[0], font_name="Times New Roman", font_size_pt=13, bold=True)

    # Title
    p_title = doc.add_paragraph()
    p_title.text = "PHIẾU XÁC ĐỊNH MỨC ĐỘ KHUYẾT TẬT ĐỐI VỚI NGƯỜI TỪ ĐỦ 6 TUỔI TRỞ LÊN"
    set_paragraph_format(p_title, WD_ALIGN_PARAGRAPH.CENTER, space_before=12, space_after=12)
    format_run(p_title.runs[0], font_name="Times New Roman", font_size_pt=14, bold=True)

    # Section I
    p_s1 = doc.add_paragraph("I. Thông tin người được xác định mức độ khuyết tật")
    set_paragraph_format(p_s1, WD_ALIGN_PARAGRAPH.LEFT, space_before=6, space_after=3)
    format_run(p_s1.runs[0], font_name="Times New Roman", font_size_pt=13, bold=True)

    lines_s1 = [
        "- Họ và tên: {{HO_TEN}}",
        "- Sinh ngày {{NGAY_SINH_NGAY}} tháng {{NGAY_SINH_THANG}} năm {{NGAY_SINH_NAM}}    Giới tính: {{GIOI_TINH}}",
        "- Hộ khẩu thường trú: Thôn {{THON}}, xã {{XA}}, huyện {{HUYEN}}, tỉnh {{TINH}}",
        "- Số CMND / CCCD: {{CMND}}",
        "- Nơi ở hiện nay: Thôn {{THON}}, xã {{XA}}, huyện {{HUYEN}}, tỉnh {{TINH}}"
    ]
    for line in lines_s1:
        p = doc.add_paragraph(line)
        set_paragraph_format(p, space_after=2)
        format_run(p.runs[0], font_name="Times New Roman", font_size_pt=13)

    # Section II
    p_s2 = doc.add_paragraph("II. Thông tin người đại diện hợp pháp (nếu có)")
    set_paragraph_format(p_s2, WD_ALIGN_PARAGRAPH.LEFT, space_before=6, space_after=3)
    format_run(p_s2.runs[0], font_name="Times New Roman", font_size_pt=13, bold=True)

    lines_s2 = [
        "- Họ và tên người đại diện: {{HO_TEN_NDH}}",
        "- Mối quan hệ với người được xác định khuyết tật: {{MOI_QUAN_HE}}",
        "- Số CMND/CCCD: {{CMND_NDH}}",
        "- Hộ khẩu thường trú: Thôn {{THON}}, xã {{XA}}, huyện {{HUYEN}}, tỉnh {{TINH}}"
    ]
    for line in lines_s2:
        p = doc.add_paragraph(line)
        set_paragraph_format(p, space_after=2)
        format_run(p.runs[0], font_name="Times New Roman", font_size_pt=13)

    # Section III
    p_s3 = doc.add_paragraph("III. Xác định dạng khuyết tật")
    set_paragraph_format(p_s3, WD_ALIGN_PARAGRAPH.LEFT, space_before=6, space_after=3)
    format_run(p_s3.runs[0], font_name="Times New Roman", font_size_pt=13, bold=True)

    p_s3_sub = doc.add_paragraph("Đã kiểm tra và đánh giá dạng khuyết tật theo Bảng tiêu chí ban hành kèm theo Thông tư 01/2019/TT-BLĐTBXH.")
    set_paragraph_format(p_s3_sub, space_after=4)
    format_run(p_s3_sub.runs[0], font_name="Times New Roman", font_size_pt=13, italic=True)

    # Section IV
    p_s4 = doc.add_paragraph("IV. Đánh giá mức độ thực hiện sinh hoạt cá nhân")
    set_paragraph_format(p_s4, WD_ALIGN_PARAGRAPH.LEFT, space_before=6, space_after=3)
    format_run(p_s4.runs[0], font_name="Times New Roman", font_size_pt=13, bold=True)

    p_score = doc.add_paragraph("- Tổng số điểm đánh giá sinh hoạt cá nhân: {{TONG_DIEM}} điểm.")
    set_paragraph_format(p_score, space_after=4)
    format_run(p_score.runs[0], font_name="Times New Roman", font_size_pt=13)

    # Section V
    p_s5 = doc.add_paragraph("V. Đề xuất xác định mức độ khuyết tật")
    set_paragraph_format(p_s5, WD_ALIGN_PARAGRAPH.LEFT, space_before=6, space_after=3)
    format_run(p_s5.runs[0], font_name="Times New Roman", font_size_pt=13, bold=True)

    lines_s5 = [
        "1. Dạng khuyết tật: {{DANG_TAT}}",
        "2. Mức độ khuyết tật: {{MUC_DO}}",
        "3. Ghi chú bổ sung: {{GHI_CHU}}"
    ]
    for line in lines_s5:
        p = doc.add_paragraph(line)
        set_paragraph_format(p, space_after=3)
        format_run(p.runs[0], font_name="Times New Roman", font_size_pt=13)

    # Signatures
    p_date = doc.add_paragraph("{{XA}}, ngày {{NGAY_HOP_NGAY}} tháng {{NGAY_HOP_THANG}} năm {{NGAY_HOP_NAM}}")
    set_paragraph_format(p_date, WD_ALIGN_PARAGRAPH.RIGHT, space_before=12, space_after=4)
    format_run(p_date.runs[0], font_name="Times New Roman", font_size_pt=13, italic=True)

    sig_table = doc.add_table(rows=2, cols=2)
    sig_table.alignment = WD_TABLE_ALIGNMENT.CENTER
    c1, c2 = sig_table.rows[0].cells

    p_c1 = c1.paragraphs[0]
    p_c1.text = "NGƯỜI GHI PHIẾU\n(Ký, ghi rõ họ tên)"
    set_paragraph_format(p_c1, WD_ALIGN_PARAGRAPH.CENTER)
    format_run(p_c1.runs[0], font_name="Times New Roman", font_size_pt=12, bold=True)

    p_c2 = c2.paragraphs[0]
    p_c2.text = "CHỦ TỊCH HỘI ĐỒNG\n(Ký tên, đóng dấu)"
    set_paragraph_format(p_c2, WD_ALIGN_PARAGRAPH.CENTER)
    format_run(p_c2.runs[0], font_name="Times New Roman", font_size_pt=12, bold=True)

    c1_b, c2_b = sig_table.rows[1].cells
    p_c1_b = c1_b.paragraphs[0]
    p_c1_b.text = "\n\n\n{{NGUOI_LAP}}"
    set_paragraph_format(p_c1_b, WD_ALIGN_PARAGRAPH.CENTER)
    format_run(p_c1_b.runs[0], font_name="Times New Roman", font_size_pt=12, bold=True)

    p_c2_b = c2_b.paragraphs[0]
    p_c2_b.text = "\n\n\n{{CHU_TICH}}"
    set_paragraph_format(p_c2_b, WD_ALIGN_PARAGRAPH.CENTER)
    format_run(p_c2_b.runs[0], font_name="Times New Roman", font_size_pt=12, bold=True)

    if data_dict:
        fill_docx_placeholders(doc, data_dict)

    return doc


def create_bien_ban_hop_doc(data_dict=None):
    """Creates the 'Biên Bản Họp Hội Đồng Kết Luận Dạng Tật và Mức Độ Khuyết Tật' document."""
    doc = docx.Document()
    section = doc.sections[0]
    set_a4_margins(section)

    # Header Table
    table = doc.add_table(rows=1, cols=2)
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    table.autofit = False

    cell_l, cell_r = table.rows[0].cells
    cell_l.width = Cm(8.5)
    cell_r.width = Cm(9.5)

    p_l1 = cell_l.paragraphs[0]
    p_l1.text = "UBND XÃ {{XA_UPPER}}"
    set_paragraph_format(p_l1, WD_ALIGN_PARAGRAPH.CENTER, space_after=1)
    format_run(p_l1.runs[0], font_name="Times New Roman", font_size_pt=12, bold=True)

    p_l2 = cell_l.add_paragraph("HỘI ĐỒNG XÁC ĐỊNH MĐKT")
    set_paragraph_format(p_l2, WD_ALIGN_PARAGRAPH.CENTER, space_after=1)
    format_run(p_l2.runs[0], font_name="Times New Roman", font_size_pt=12, bold=True)

    p_r1 = cell_r.paragraphs[0]
    p_r1.text = "CỘNG HÒA XÃ HỘI CHỦ NGHĨA VIỆT NAM"
    set_paragraph_format(p_r1, WD_ALIGN_PARAGRAPH.CENTER, space_after=1)
    format_run(p_r1.runs[0], font_name="Times New Roman", font_size_pt=12, bold=True)

    p_r2 = cell_r.add_paragraph("Độc lập – Tự do – Hạnh phúc")
    set_paragraph_format(p_r2, WD_ALIGN_PARAGRAPH.CENTER, space_after=6)
    format_run(p_r2.runs[0], font_name="Times New Roman", font_size_pt=13, bold=True)

    # Title
    p_title1 = doc.add_paragraph("BIÊN BẢN")
    set_paragraph_format(p_title1, WD_ALIGN_PARAGRAPH.CENTER, space_before=10, space_after=2)
    format_run(p_title1.runs[0], font_name="Times New Roman", font_size_pt=14, bold=True)

    p_title2 = doc.add_paragraph("HỌP KẾT LUẬN DẠNG KHUYẾT TẬT VÀ MỨC ĐỘ KHUYẾT TẬT")
    set_paragraph_format(p_title2, WD_ALIGN_PARAGRAPH.CENTER, space_after=10)
    format_run(p_title2.runs[0], font_name="Times New Roman", font_size_pt=14, bold=True)

    # I. Thời gian, địa điểm
    p_i = doc.add_paragraph("I. Thời gian, địa điểm")
    set_paragraph_format(p_i, space_before=4, space_after=2)
    format_run(p_i.runs[0], font_name="Times New Roman", font_size_pt=13, bold=True)

    p_i_detail = doc.add_paragraph("Hôm nay, vào hồi {{GIO_HOP}} giờ {{PHUT_HOP}} phút ngày {{NGAY_HOP_NGAY}} tháng {{NGAY_HOP_THANG}} năm {{NGAY_HOP_NAM}}, Tại {{DIA_DIEM_HOP}}.")
    set_paragraph_format(p_i_detail, space_after=4)
    format_run(p_i_detail.runs[0], font_name="Times New Roman", font_size_pt=13)

    # II. Thành phần
    p_ii = doc.add_paragraph("II. Thành phần Hội đồng xác định mức độ khuyết tật")
    set_paragraph_format(p_ii, space_before=4, space_after=2)
    format_run(p_ii.runs[0], font_name="Times New Roman", font_size_pt=13, bold=True)

    members = [
        "1. Ông (Bà): {{CHU_TICH}} - Chủ tịch Hội đồng, chủ trì cuộc họp;",
        "2. Ông (Bà): {{Y_TE}} - Trưởng Trạm y tế cấp xã, thành viên;",
        "3. Ông (Bà): {{MTTQ}} - Đại diện UB MTTQ xã, thành viên;",
        "4. Ông (Bà): {{NGUOI_GHI_BIEN_BAN}} - Công chức Văn hóa - Xã hội xã, người ghi biên bản."
    ]
    for m in members:
        p = doc.add_paragraph(m)
        set_paragraph_format(p, space_after=2)
        format_run(p.runs[0], font_name="Times New Roman", font_size_pt=13)

    # III. Nội dung
    p_iii = doc.add_paragraph("III. Nội dung")
    set_paragraph_format(p_iii, space_before=4, space_after=2)
    format_run(p_iii.runs[0], font_name="Times New Roman", font_size_pt=13, bold=True)

    p_p1 = doc.add_paragraph("1. Xác định dạng khuyết tật và mức độ khuyết tật cho:")
    set_paragraph_format(p_p1, space_after=2)
    format_run(p_p1.runs[0], font_name="Times New Roman", font_size_pt=13, bold=True)

    lines_obj = [
        "   - Ông (bà): {{HO_TEN}}    Giới tính: {{GIOI_TINH}}",
        "   - Ngày, tháng, năm sinh: {{NGAY_SINH}}",
        "   - Nơi ở hiện nay: Thôn {{THON}}, xã {{XA}}, huyện {{HUYEN}}, tỉnh {{TINH}}"
    ]
    for l in lines_obj:
        p = doc.add_paragraph(l)
        set_paragraph_format(p, space_after=2)
        format_run(p.runs[0], font_name="Times New Roman", font_size_pt=13)

    p_p2 = doc.add_paragraph("2. Hội đồng quan sát, phỏng vấn người được xác định mức độ khuyết tật hoặc người đại diện hợp pháp của họ.")
    set_paragraph_format(p_p2, space_after=2)
    format_run(p_p2.runs[0], font_name="Times New Roman", font_size_pt=13)

    p_p3 = doc.add_paragraph("3. Công chức Văn hóa - Xã hội cấp xã báo cáo kết quả thu thập thông tin Phiếu xác định mức độ khuyết tật.")
    set_paragraph_format(p_p3, space_after=2)
    format_run(p_p3.runs[0], font_name="Times New Roman", font_size_pt=13)

    p_p4 = doc.add_paragraph("4. Ý kiến của các thành viên dự họp: Thống nhất hoàn toàn với kết quả thu thập thông tin và đánh giá thực tế đối tượng.")
    set_paragraph_format(p_p4, space_after=2)
    format_run(p_p4.runs[0], font_name="Times New Roman", font_size_pt=13)

    p_p5 = doc.add_paragraph("5. Kết quả biểu quyết: 100% thành viên Hội đồng nhất trí.")
    set_paragraph_format(p_p5, space_after=2)
    format_run(p_p5.runs[0], font_name="Times New Roman", font_size_pt=13)

    p_p6 = doc.add_paragraph("6. Kết luận: Hội đồng thống nhất kết luận như sau:")
    set_paragraph_format(p_p6, space_after=2)
    format_run(p_p6.runs[0], font_name="Times New Roman", font_size_pt=13, bold=True)

    p_conc = doc.add_paragraph("   - Dạng khuyết tật: {{DANG_TAT}}\n   - Mức độ khuyết tật: {{MUC_DO}}\n   - Ghi chú: {{GHI_CHU}}")
    set_paragraph_format(p_conc, space_after=6)
    format_run(p_conc.runs[0], font_name="Times New Roman", font_size_pt=13)

    p_end = doc.add_paragraph("Cuộc họp kết thúc hồi {{GIO_KET_THUC}} giờ ngày {{NGAY_HOP_NGAY}} tháng {{NGAY_HOP_THANG}} năm {{NGAY_HOP_NAM}}.\nBiên bản này được lập thành 03 bản có giá trị pháp lý như nhau./.")
    set_paragraph_format(p_end, space_after=12)
    format_run(p_end.runs[0], font_name="Times New Roman", font_size_pt=13, italic=True)

    # Signatures
    sig_table = doc.add_table(rows=2, cols=2)
    sig_table.alignment = WD_TABLE_ALIGNMENT.CENTER
    c1, c2 = sig_table.rows[0].cells

    p_c1 = c1.paragraphs[0]
    p_c1.text = "NGƯỜI GHI BIÊN BẢN"
    set_paragraph_format(p_c1, WD_ALIGN_PARAGRAPH.CENTER)
    format_run(p_c1.runs[0], font_name="Times New Roman", font_size_pt=12, bold=True)

    p_c2 = c2.paragraphs[0]
    p_c2.text = "CHỦ TỊCH HỘI ĐỒNG"
    set_paragraph_format(p_c2, WD_ALIGN_PARAGRAPH.CENTER)
    format_run(p_c2.runs[0], font_name="Times New Roman", font_size_pt=12, bold=True)

    c1_b, c2_b = sig_table.rows[1].cells
    p_c1_b = c1_b.paragraphs[0]
    p_c1_b.text = "\n\n\n{{NGUOI_GHI_BIEN_BAN}}"
    set_paragraph_format(p_c1_b, WD_ALIGN_PARAGRAPH.CENTER)
    format_run(p_c1_b.runs[0], font_name="Times New Roman", font_size_pt=12, bold=True)

    p_c2_b = c2_b.paragraphs[0]
    p_c2_b.text = "\n\n\n{{CHU_TICH}}"
    set_paragraph_format(p_c2_b, WD_ALIGN_PARAGRAPH.CENTER)
    format_run(p_c2_b.runs[0], font_name="Times New Roman", font_size_pt=12, bold=True)

    if data_dict:
        fill_docx_placeholders(doc, data_dict)

    return doc


def create_thong_bao_niem_yet_doc(data_dict=None, df_list=None):
    """Creates the 'Thông Báo Niêm Yết Kết Quả Xác Định Mức Độ Khuyết Tật' document with table list in A4 Landscape."""
    doc = docx.Document()
    
    # --- SECTION 1: Notice Document (A4 Portrait) ---
    section1 = doc.sections[0]
    set_a4_margins(section1)

    # Header Table
    table = doc.add_table(rows=1, cols=2)
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    table.autofit = False

    cell_l, cell_r = table.rows[0].cells
    cell_l.width = Cm(7.5)
    cell_r.width = Cm(9.0)

    p_l1 = cell_l.paragraphs[0]
    p_l1.text = "ỦY BAN NHÂN DÂN"
    set_paragraph_format(p_l1, WD_ALIGN_PARAGRAPH.CENTER, space_after=1)
    format_run(p_l1.runs[0], font_name="Times New Roman", font_size_pt=12, bold=True)

    p_l2 = cell_l.add_paragraph("XÃ {{XA_UPPER}}")
    set_paragraph_format(p_l2, WD_ALIGN_PARAGRAPH.CENTER, space_after=1)
    format_run(p_l2.runs[0], font_name="Times New Roman", font_size_pt=12, bold=True)

    p_l3 = cell_l.add_paragraph("Số: {{SO_THONG_BAO}}")
    set_paragraph_format(p_l3, WD_ALIGN_PARAGRAPH.CENTER, space_after=6)
    format_run(p_l3.runs[0], font_name="Times New Roman", font_size_pt=12)

    p_r1 = cell_r.paragraphs[0]
    p_r1.text = "CỘNG HÒA XÃ HỘI CHỦ NGHĨA VIỆT NAM"
    set_paragraph_format(p_r1, WD_ALIGN_PARAGRAPH.CENTER, space_after=1)
    format_run(p_r1.runs[0], font_name="Times New Roman", font_size_pt=12, bold=True)

    p_r2 = cell_r.add_paragraph("Độc lập - Tự do - Hạnh phúc")
    set_paragraph_format(p_r2, WD_ALIGN_PARAGRAPH.CENTER, space_after=1)
    format_run(p_r2.runs[0], font_name="Times New Roman", font_size_pt=13, bold=True)

    p_r3 = cell_r.add_paragraph("{{XA}}, ngày {{NGAY_HOP_NGAY}} tháng {{NGAY_HOP_THANG}} năm {{NGAY_HOP_NAM}}")
    set_paragraph_format(p_r3, WD_ALIGN_PARAGRAPH.CENTER, space_after=6)
    format_run(p_r3.runs[0], font_name="Times New Roman", font_size_pt=13, italic=True)

    # Title
    p_t1 = doc.add_paragraph("THÔNG BÁO")
    set_paragraph_format(p_t1, WD_ALIGN_PARAGRAPH.CENTER, space_before=10, space_after=2)
    format_run(p_t1.runs[0], font_name="Times New Roman", font_size_pt=14, bold=True)

    p_t2 = doc.add_paragraph("Về việc niêm yết công khai kết quả họp kết luận")
    set_paragraph_format(p_t2, WD_ALIGN_PARAGRAPH.CENTER, space_after=1)
    format_run(p_t2.runs[0], font_name="Times New Roman", font_size_pt=14, bold=True)

    p_t3 = doc.add_paragraph("dạng tật và mức độ khuyết tật")
    set_paragraph_format(p_t3, WD_ALIGN_PARAGRAPH.CENTER, space_after=10)
    format_run(p_t3.runs[0], font_name="Times New Roman", font_size_pt=14, bold=True)

    # Căn cứ (Decrees / Legal basis) - Exact standard formatting with justification and indentation
    d1 = "Căn cứ Nghị định số 28/2012/NĐ-CP ngày 10/4/2012 của Chính phủ quy định chi tiết và hướng dẫn thi hành một số điều của Luật Người khuyết tật;"
    p_d1 = doc.add_paragraph(d1)
    set_paragraph_format(p_d1, WD_ALIGN_PARAGRAPH.JUSTIFY, space_after=3, line_spacing=1.15, first_line_indent=28)
    format_run(p_d1.runs[0], font_name="Times New Roman", font_size_pt=13)

    d2 = "Căn cứ Thông tư số 01/2019/TT-BLĐTBXH ngày 02/01/2019 của Bộ Lao động - Thương binh và Xã hội quy định về việc xác định mức độ khuyết tật do Hội đồng xác định mức độ khuyết tật thực hiện;"
    p_d2 = doc.add_paragraph(d2)
    set_paragraph_format(p_d2, WD_ALIGN_PARAGRAPH.JUSTIFY, space_after=3, line_spacing=1.15, first_line_indent=28)
    format_run(p_d2.runs[0], font_name="Times New Roman", font_size_pt=13)

    d3 = "Căn cứ Thông tư số 19/2026/TT-BYT ngày 09/6/2026 của Bộ Y tế sửa đổi, bổ sung một số điều của Thông tư số 01/2019/TT-BLĐTBXH ngày 02 tháng 01 năm 2019 của Bộ trưởng Bộ Lao động - Thương binh và Xã hội quy định về việc xác định mức độ khuyết tật do Hội đồng xác định mức độ khuyết tật thực hiện;"
    p_d3 = doc.add_paragraph(d3)
    set_paragraph_format(p_d3, WD_ALIGN_PARAGRAPH.JUSTIFY, space_after=3, line_spacing=1.15, first_line_indent=28)
    format_run(p_d3.runs[0], font_name="Times New Roman", font_size_pt=13)

    d4 = "Căn cứ kết quả họp kết luận dạng khuyết tật và mức độ khuyết tật ngày {{NGAY_HOP_NGAY}}/{{NGAY_HOP_THANG}}/{{NGAY_HOP_NAM}} của Hội đồng xác định mức độ khuyết tật xã {{XA}}."
    p_d4 = doc.add_paragraph(d4)
    set_paragraph_format(p_d4, WD_ALIGN_PARAGRAPH.JUSTIFY, space_after=6, line_spacing=1.15, first_line_indent=28)
    format_run(p_d4.runs[0], font_name="Times New Roman", font_size_pt=13)

    # Body Paragraph 1
    p_b1 = doc.add_paragraph()
    set_paragraph_format(p_b1, WD_ALIGN_PARAGRAPH.JUSTIFY, space_after=3, line_spacing=1.15, first_line_indent=28)
    r1 = p_b1.add_run("Ủy ban nhân dân xã {{XA}} tổ chức tiến hành niêm yết công khai kết quả họp kết luận dạng khuyết tật và mức độ khuyết tật cho {{SO_LUONG_TRUONG_HOP}} trường hợp ")
    format_run(r1, font_name="Times New Roman", font_size_pt=13)
    r2 = p_b1.add_run("(Có danh sách kèm theo)")
    format_run(r2, font_name="Times New Roman", font_size_pt=13, italic=True)

    # Sub-bullets (Indented, without minus bullets)
    p_b2 = doc.add_paragraph("Địa điểm: {{DIA_DIEM_NIEM_YET}}")
    set_paragraph_format(p_b2, WD_ALIGN_PARAGRAPH.JUSTIFY, space_after=2, line_spacing=1.15, first_line_indent=48)
    format_run(p_b2.runs[0], font_name="Times New Roman", font_size_pt=13)

    p_b3 = doc.add_paragraph("Thời gian niêm yết: {{THOI_GIAN_NIEM_YET}} kể từ ngày niêm yết.")
    set_paragraph_format(p_b3, WD_ALIGN_PARAGRAPH.JUSTIFY, space_after=4, line_spacing=1.15, first_line_indent=48)
    format_run(p_b3.runs[0], font_name="Times New Roman", font_size_pt=13)

    # Body Paragraph 2
    p_b4 = doc.add_paragraph("Ủy ban nhân dân xã {{XA}} thông báo để các cá nhân, hộ gia đình nếu có thắc mắc, kiến nghị về nội dung công khai trên thì nộp đơn tại trụ sở Ủy ban nhân dân xã {{XA}} để được giải quyết. Sau thời gian niêm yết, nếu không có ý kiến khiếu nại của người dân thì Ủy ban nhân dân xã sẽ tiến hành làm các thủ tục tiếp theo cho đối tượng trên theo quy định./.")
    set_paragraph_format(p_b4, WD_ALIGN_PARAGRAPH.JUSTIFY, space_after=12, line_spacing=1.15, first_line_indent=28)
    format_run(p_b4.runs[0], font_name="Times New Roman", font_size_pt=13)

    # Signatures Table
    sig_tbl = doc.add_table(rows=2, cols=2)
    sig_tbl.alignment = WD_TABLE_ALIGNMENT.CENTER
    c1, c2 = sig_tbl.rows[0].cells

    p_c1 = c1.paragraphs[0]
    set_paragraph_format(p_c1, WD_ALIGN_PARAGRAPH.LEFT, line_spacing=1.1)
    r_nr = p_c1.add_run("Nơi nhận:\n")
    format_run(r_nr, font_name="Times New Roman", font_size_pt=11, bold=True, italic=True)
    r_lst = p_c1.add_run("- Niêm yết tại Trụ sở UBND xã;\n- Lưu: VT, UBND.")
    format_run(r_lst, font_name="Times New Roman", font_size_pt=11, italic=True)

    p_c2 = c2.paragraphs[0]
    p_c2.text = "KT. CHỦ TỊCH\nPHÓ CHỦ TỊCH"
    set_paragraph_format(p_c2, WD_ALIGN_PARAGRAPH.CENTER, line_spacing=1.1)
    format_run(p_c2.runs[0], font_name="Times New Roman", font_size_pt=12, bold=True)

    c1_b, c2_b = sig_tbl.rows[1].cells
    p_c2_b = c2_b.paragraphs[0]
    p_c2_b.text = "\n\n\n{{CHU_TICH}}"
    set_paragraph_format(p_c2_b, WD_ALIGN_PARAGRAPH.CENTER)
    format_run(p_c2_b.runs[0], font_name="Times New Roman", font_size_pt=12, bold=True)

    # --- SECTION 2: List Table (A4 LANDSCAPE) ---
    section2 = doc.add_section(docx.enum.section.WD_SECTION.NEW_PAGE)
    section2.orientation = docx.enum.section.WD_ORIENT.LANDSCAPE
    section2.page_width = Cm(29.7)
    section2.page_height = Cm(21.0)
    section2.top_margin = Cm(2.0)
    section2.bottom_margin = Cm(2.0)
    section2.left_margin = Cm(2.0)
    section2.right_margin = Cm(2.0)

    # Page 2 Title
    p_p2_t1 = doc.add_paragraph("DANH SÁCH THÔNG BÁO NIÊM YẾT KẾT QUẢ XÁC ĐỊNH MỨC ĐỘ KHUYẾT TẬT")
    set_paragraph_format(p_p2_t1, WD_ALIGN_PARAGRAPH.CENTER, space_before=10, space_after=2)
    format_run(p_p2_t1.runs[0], font_name="Times New Roman", font_size_pt=13, bold=True)

    p_p2_sub = doc.add_paragraph("(Kèm theo Thông báo số {{SO_THONG_BAO}} ngày {{NGAY_HOP_NGAY}}/{{NGAY_HOP_THANG}}/{{NGAY_HOP_NAM}} của Ủy ban nhân dân xã {{XA}})")
    set_paragraph_format(p_p2_sub, WD_ALIGN_PARAGRAPH.CENTER, space_after=10)
    format_run(p_p2_sub.runs[0], font_name="Times New Roman", font_size_pt=12, italic=True)

    # Table in Landscape
    list_table = doc.add_table(rows=1, cols=7)
    list_table.alignment = WD_TABLE_ALIGNMENT.CENTER
    list_table.autofit = False

    headers = ["TT", "Họ và tên", "Ngày tháng năm sinh", "Thôn", "Dạng tật", "Mức độ khuyết tật", "Ghi chú"]
    widths = [Cm(1.2), Cm(5.0), Cm(3.2), Cm(3.8), Cm(4.0), Cm(4.0), Cm(4.5)]

    hdr_cells = list_table.rows[0].cells
    for idx, text in enumerate(headers):
        hdr_cells[idx].width = widths[idx]
        hdr_cells[idx].vertical_alignment = WD_ALIGN_VERTICAL.CENTER
        p = hdr_cells[idx].paragraphs[0]
        p.text = text
        set_paragraph_format(p, WD_ALIGN_PARAGRAPH.CENTER, space_before=3, space_after=3)
        format_run(p.runs[0], font_name="Times New Roman", font_size_pt=11, bold=True)

    set_table_borders(list_table)

    if data_dict:
        fill_docx_placeholders(doc, data_dict)

    if df_list:
        append_table_rows(doc, df_list)

    return doc
